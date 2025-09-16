import pandas as pd
import numpy as np
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows

processor_names = [
    'Ryzen 5 3600', 'Ryzen 7 5800X', 'Ryzen 9 5900X', 'Core i5-12400F',
    'Core i7-12700K', 'Core i9-13900K', 'Ryzen 5 5600X', 'Ryzen 7 3700X',
    'Core i5-11400F', 'Core i3-10100'
]

np.random.seed(42)
data = pd.DataFrame({
    'Имя': processor_names[:10],
    'Количество': np.random.randint(1, 200, size=10),
    'Цена': np.round(np.random.uniform(20000, 50000, size=10), 2)
})

print(data)

data.to_excel("output.xlsx", index=False, sheet_name='List1')

wb = load_workbook("output.xlsx")
ws = wb['List1']

red_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")

header = [cell.value for cell in ws[1]]
col_idx = header.index('Количество') + 1

for row in range(2, ws.max_row + 1):
    cell = ws.cell(row=row, column=col_idx)
    if isinstance(cell.value, (int, float)) and cell.value > 100:
        cell.fill = red_fill

wb.save("beatiful_output.xlsx")

from docx import Document

document = Document()

document.add_heading('Отчет')
document.add_paragraph('Отчет сгенерирован роботом...')

table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Имя'
hdr_cells[1].text = 'Количество'
hdr_cells[2].text = 'Цена'

for index, row in data.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Имя'])
    row_cells[1].text = str(row['Количество'])
    row_cells[2].text = str(row['Цена'])

document.save('output.docx')

from fpdf import FPDF
import csv

pdf = FPDF()
pdf.add_page()
pdf.set_font('helvetica', size=13)

excel_data = pd.read_excel('beatiful_output.xlsx')

with pdf.table() as table:
    for _, data_row in excel_data.iterrows():
        row = table.row()
        row.cell(str(data_row['Имя']))
        row.cell(str(data_row['Количество']))
        row.cell(str(data_row['Цена']))

pdf.output("output.pdf")

from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO

def create_overlay(text_data):
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.setFont("Helvetica", 13)

    for text, x, y in text_data:
        can.drawString(x, y, text)

    can.save()
    packet.seek(0)
    return packet

overlay_data = [
    ("16.09.2025", 400, 500),
    ("150 000", 300, 500),
]

overlay_pdf = create_overlay(overlay_data)

reader = PdfReader("output.pdf")
page = reader.pages[0]

overlay_reader = PdfReader(overlay_pdf)
overlay_page = overlay_reader.pages[0]

page.merge_page(overlay_page)

writer = PdfWriter()
writer.add_page(page)

with open("filled_form.pdf", "wb") as output_file:
    writer.write(output_file)